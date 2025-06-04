from docx import Document
from docx.shared import Pt

# Criar o documento
doc = Document()

# Função para adicionar título formatado
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Função para adicionar parágrafo com espaçamento
def add_paragraph(text):
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(11)

# Adicionar informações ao currículo
add_heading("HUGO CASTRO SILVA", 0)
add_paragraph("📞 (11) 96858-3746")
add_paragraph("📧 hugocastrosilva679@gmail.com")
add_paragraph("🌐 https://www.linkedin.com/in/hugo-castro-1b4761241/")
add_paragraph("📍 São Paulo – SP")

add_heading("OBJETIVO PROFISSIONAL", 1)
add_paragraph("Busco meu primeiro estágio na área de Análise e Desenvolvimento de Sistemas, com o objetivo de aplicar meus conhecimentos teóricos e desenvolver habilidades práticas no setor de tecnologia.")

add_heading("EXPERIÊNCIA PROFISSIONAL", 1)
add_paragraph("Clínica Dra. Dinorah – Atendente")
add_paragraph("📅 Novembro/2022 – Atual")
add_paragraph("- Atendimento ao público e suporte aos pacientes.\n"
              "- Organização de agendas e apoio administrativo.\n"
              "- Desenvolvimento de habilidades interpessoais e de trabalho em equipe.\n"
              "- Destaque pelo bom relacionamento com clientes e colegas de trabalho.")

add_heading("FORMAÇÃO ACADÊMICA", 1)
add_paragraph("Curso Superior em Análise e Desenvolvimento de Sistemas – UNINTER")
add_paragraph("📅 Previsão de Conclusão: Agosto/2026")

add_heading("HABILIDADES E COMPETÊNCIAS", 1)
add_paragraph("Técnicas:")
add_paragraph("- Conhecimento em MySQL básico\n"
              "- Conhecimentos em Git e GitHub\n"
              "- Aprendizado contínuo em HTML e CSS\n"
              "- Interesse por desenvolvimento web e bancos de dados")

add_paragraph("Comportamentais:")
add_paragraph("- Facilidade para trabalhar em equipe\n"
              "- Comunicação clara e objetiva\n"
              "- Proatividade e dedicação\n"
              "- Facilidade de aprendizado e adaptação")

add_heading("CURSOS E CERTIFICAÇÕES", 1)
add_paragraph("- MySQL Básico – 40 horas\n"
              "- Git e GitHub – 20 horas\n"
              "- Em andamento: HTML e CSS (estudos próprios)")

# Salvar o documento
file_path = "/mnt/data/Curriculo_Hugo_Castro_Silva.docx"
doc.save(file_path)

file_path

from docx2pdf import convert

# Converte o arquivo gerado em PDF
convert("/mnt/data/Curriculo_Hugo_Castro_Silva.docx")
